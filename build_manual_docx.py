"""
TideBedPy Word Manual Builder
전문적인 .docx 매뉴얼 자동 생성
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── 색상 팔레트 ──────────────────────────────────────────────
NAVY      = RGBColor(0x1B, 0x3A, 0x5C)
BLUE      = RGBColor(0x4A, 0x6F, 0xA5)
LIGHT_BLUE= RGBColor(0xD6, 0xE4, 0xF0)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY  = RGBColor(0x66, 0x66, 0x66)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT    = RGBColor(0x2E, 0x75, 0xB6)

FONT_KR = '맑은 고딕'
FONT_EN = 'Segoe UI'
FONT_CODE = 'Consolas'

OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        'tidebedpy', 'manual', 'TideBedPy_Manual.docx')


def set_cell_shading(cell, color_hex):
    """셀 배경색 설정."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, font_size=9, color=DARK_GRAY,
                  font_name=FONT_KR, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """셀 텍스트 설정 + 서식."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.bold = bold
    # 셀 여백
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="40" w:type="dxa"/>'
        f'  <w:left w:w="80" w:type="dxa"/>'
        f'  <w:bottom w:w="40" w:type="dxa"/>'
        f'  <w:right w:w="80" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def add_styled_table(doc, headers, rows, col_widths=None, header_color='1B3A5C'):
    """서식이 적용된 테이블 생성."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_color)
        set_cell_text(cell, h, bold=True, font_size=9, color=WHITE,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 데이터
    for r_idx, row_data in enumerate(rows):
        bg = 'F7F9FC' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_shading(cell, bg)
            set_cell_text(cell, str(val), font_size=9)

    # 열 너비
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def add_info_box(doc, title, lines, color_hex='E8F0FE'):
    """정보 박스 (1열 테이블로 구현)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, color_hex)

    # 타이틀
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.font.name = FONT_KR
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run.bold = True

    # 내용
    for line in lines:
        p2 = cell.add_paragraph()
        run2 = p2.add_run(line)
        run2.font.name = FONT_KR
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
        run2.font.size = Pt(9)
        run2.font.color.rgb = DARK_GRAY

    # 셀 여백
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="100" w:type="dxa"/>'
        f'  <w:left w:w="160" w:type="dxa"/>'
        f'  <w:bottom w:w="100" w:type="dxa"/>'
        f'  <w:right w:w="160" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

    doc.add_paragraph()  # 간격


def setup_styles(doc):
    """문서 스타일 설정."""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_KR
    font.size = Pt(10)
    font.color.rgb = DARK_GRAY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = FONT_KR
    h1.font.size = Pt(20)
    h1.font.color.rgb = NAVY
    h1.font.bold = True
    h1.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = FONT_KR
    h2.font.size = Pt(14)
    h2.font.color.rgb = BLUE
    h2.font.bold = True
    h2.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = FONT_KR
    h3.font.size = Pt(12)
    h3.font.color.rgb = ACCENT
    h3.font.bold = True
    h3.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)


def add_cover(doc):
    """표지 페이지."""
    for _ in range(4):
        doc.add_paragraph()

    # 프로그램명
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TideBedPy')
    run.font.name = FONT_KR
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
    run.font.size = Pt(42)
    run.font.color.rgb = NAVY
    run.bold = True

    # 구분선
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('━' * 40)
    run2.font.color.rgb = BLUE
    run2.font.size = Pt(12)

    # 부제
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('Co-tidal 격자 기반 조석보정 프로그램')
    run3.font.name = FONT_KR
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
    run3.font.size = Pt(18)
    run3.font.color.rgb = BLUE

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('사용자 매뉴얼')
    run4.font.name = FONT_KR
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
    run4.font.size = Pt(16)
    run4.font.color.rgb = BLUE

    doc.add_paragraph()

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run('Version 2.3.0')
    run5.font.name = FONT_KR
    run5._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
    run5.font.size = Pt(14)
    run5.font.color.rgb = NAVY
    run5.bold = True

    for _ in range(6):
        doc.add_paragraph()

    # 저작권 정보
    for text in [
        'Original : TideBedLite v1.05  (c) 2014 KHOA / GeoSR Inc.',
        'Python Rewrite : Junhyub, 2025',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = FONT_KR
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
        run.font.size = Pt(10)
        run.font.color.rgb = MID_GRAY

    doc.add_page_break()


def add_toc(doc):
    """목차 페이지."""
    doc.add_heading('목  차', level=1)
    doc.add_paragraph()

    toc_items = [
        ('1.', '프로그램 개요'),
        ('2.', '시스템 요구사항'),
        ('3.', '데이터 준비'),
        ('    3.1', '조위 데이터 (KHOA 바다누리)'),
        ('    3.2', '항적 데이터 (CARIS HIPS)'),
        ('    3.3', '기준항정보.txt'),
        ('    3.4', '표준개정수 DB (Co-tidal)'),
        ('4.', '사용 방법'),
        ('    4.1', 'GUI 모드'),
        ('    4.2', 'CLI 모드'),
        ('    4.3', '검증 모드'),
        ('5.', '출력 파일 설명'),
        ('6.', '핵심 알고리즘 및 공식'),
        ('    6.1', 'Vincenty 측지 거리'),
        ('    6.2', 'IDW 가중평균'),
        ('    6.3', 'Co-tidal 이중선형 보간'),
        ('    6.4', 'Akima 스플라인 보간'),
        ('    6.5', '조석보정 계산 단계'),
        ('    6.6', '시간대 처리'),
        ('7.', '시간대 설정 안내'),
        ('8.', 'TideBedLite 대비 개선점'),
        ('9.', '문제 해결 (FAQ)'),
        ('10.', '라이선스 및 저작권'),
    ]

    for num, title in toc_items:
        p = doc.add_paragraph()
        is_sub = num.startswith('    ')
        indent = '      ' if is_sub else ''

        run_num = p.add_run(f'{indent}{num}  ')
        run_num.font.name = FONT_KR
        run_num._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
        run_num.font.size = Pt(11 if not is_sub else 10)
        run_num.font.color.rgb = NAVY if not is_sub else BLUE
        run_num.bold = not is_sub

        run_title = p.add_run(title)
        run_title.font.name = FONT_KR
        run_title._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
        run_title.font.size = Pt(11 if not is_sub else 10)
        run_title.font.color.rgb = DARK_GRAY

    doc.add_page_break()


def add_section_1(doc):
    """1. 프로그램 개요."""
    doc.add_heading('1. 프로그램 개요', level=1)

    # 기본 정보 테이블
    add_styled_table(doc,
        ['항목', '내용'],
        [
            ['프로그램명', 'TideBedPy'],
            ['원본 프로그램', 'TideBedLite v1.05 (c) 2014, KHOA / GeoSR Inc.'],
            ['Python 재작성', 'Junhyub, 2025'],
            ['현재 버전', 'v2.3.0'],
        ],
        col_widths=[4, 12]
    )
    doc.add_paragraph()

    doc.add_heading('목적', level=2)
    doc.add_paragraph(
        'TideBedPy는 해양 수로측량에서 수집된 항적(Navigation) 데이터에 대해 '
        'Co-tidal 격자 기반 조석보정값(Tc)을 산출하는 전문 프로그램입니다.'
    )
    doc.add_paragraph(
        '수로측량 시 측심기(Echo Sounder)로 취득한 수심 데이터는 조석의 영향을 '
        '받으므로, 정확한 해저 지형도 작성을 위해 조석보정이 필수입니다.'
    )

    doc.add_heading('주요 기능', level=2)
    features = [
        '복수 기준항(조위관측소)의 실측/예측 조위 데이터 로드',
        '각 항적 포인트에서 Co-tidal 격자를 이용한 HRatio·TimeCorrector 계산',
        'IDW(역거리 가중) 보간법으로 복수 기준항의 조석보정값 가중평균',
        'CARIS HIPS 호환 .tid 형식 결과 출력',
    ]
    for f in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f)
        run.font.name = FONT_KR
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
        run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph(
        '본 프로그램은 KHOA/GeoSR의 TideBedLite v1.05(.NET)를 Python으로 완전히 '
        '재구현한 것으로, 원본의 핵심 알고리즘을 충실히 재현하면서 다양한 Nav 포맷 '
        '자동 감지, 유연한 시간대 설정, Akima 스플라인, 시각화 등을 추가하여 '
        '사용성과 정확도를 개선하였습니다.'
    )
    doc.add_page_break()


def add_section_2(doc):
    """2. 시스템 요구사항."""
    doc.add_heading('2. 시스템 요구사항', level=1)

    doc.add_heading('운영체제', level=2)
    doc.add_paragraph('• Windows 10 (64-bit) 이상')
    doc.add_paragraph('• Windows 11 (64-bit) 지원')

    doc.add_heading('실행 환경', level=2)

    add_info_box(doc, '방법 1 — Python 환경', [
        'Python 3.8 이상 (3.10+ 권장)',
        '',
        '필수 패키지:',
        '  • numpy — 수치 연산 및 배열 처리',
        '  • scipy — Akima 스플라인 보간',
        '  • netCDF4 — Co-tidal NetCDF 격자 파일 읽기',
        '  • geographiclib — Vincenty/Karney 측지 거리 계산',
        '  • chardet — 파일 인코딩 자동 감지',
        '  • matplotlib — 그래프 및 지도 시각화',
        '',
        '설치: pip install numpy scipy netCDF4 geographiclib chardet matplotlib',
    ], 'E8F0FE')

    add_info_box(doc, '방법 2 — EXE 독립 실행', [
        'TideBedPy.exe 더블클릭으로 실행 (Python 설치 불필요)',
        '모든 의존성이 EXE 내부에 포함',
    ], 'E8F8E8')

    doc.add_heading('하드웨어 권장사양', level=2)
    add_styled_table(doc,
        ['항목', '권장 사양'],
        [
            ['CPU', 'Intel Core i5 이상 (멀티코어 권장)'],
            ['RAM', '4GB 이상 (대규모 항적 시 8GB 권장)'],
            ['디스크', '500MB 이상 여유 공간'],
            ['모니터', '1280×720 이상 (GUI 모드)'],
        ],
        col_widths=[4, 12]
    )
    doc.add_page_break()


def add_section_3(doc):
    """3. 데이터 준비."""
    doc.add_heading('3. 데이터 준비', level=1)

    doc.add_paragraph('TideBedPy 실행에 필요한 4종 데이터:')
    add_styled_table(doc,
        ['#', '데이터', '설명'],
        [
            ['1', '조위 데이터 (TOPS)', '기준항별 실측 또는 예측 조위'],
            ['2', '항적 데이터 (Nav)', 'CARIS HIPS에서 Export한 항적'],
            ['3', '기준항정보.txt', '기준항 목록 및 조화상수'],
            ['4', '표준개정수 DB', 'Co-tidal NetCDF 격자 파일'],
        ],
        col_widths=[1.5, 5, 9.5]
    )
    doc.add_paragraph()

    # 3.1
    doc.add_heading('3.1  조위 데이터 취득 (KHOA 바다누리)', level=2)
    doc.add_paragraph('접속: https://www.khoa.go.kr/oceangrid/gis/category/observe/observeSearch.do')
    doc.add_paragraph()
    doc.add_heading('취득 절차', level=3)
    steps = [
        '바다누리 해양정보서비스 접속',
        '조석/조류 → 조석관측 선택',
        '원하는 관측소(기준항) 선택',
        '조회 기간 설정 (측량 기간 전후 1일 이상 여유)',
        '데이터 유형 선택 (실측조위 권장)',
        'TOPS 형식으로 다운로드',
        '동일 폴더에 모아서 저장',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'{i}) {s}')

    doc.add_paragraph()
    doc.add_heading('TOPS 파일 형식', level=3)
    add_info_box(doc, 'TOPS 파일 구조', [
        '인천                     ← 관측소 명칭 (1행)',
        'cm                       ← 단위 (2행)',
        '2024 01 15 00 00  345    ← 데이터 행',
        '2024 01 15 00 10  342',
        '...',
        '',
        '데이터 행 형식: YYYY MM DD HH MM LevelCm',
    ], 'FFF8E1')

    add_info_box(doc, '⚠ 주의사항', [
        '• 기준항정보.txt의 Name 필드와 TOPS 1행의 관측소명이 정확히 일치해야 함',
        '• 여러 기준항 파일을 동일 폴더에 저장',
        '• EUC-KR / UTF-8 인코딩 자동 감지',
    ], 'FDE8E8')

    # 3.2
    doc.add_heading('3.2  항적(Navigation) 데이터 (CARIS HIPS)', level=2)
    doc.add_heading('CARIS HIPS Export 절차', level=3)
    caris_steps = [
        'CARIS HIPS & SIPS 실행',
        '보정 대상 측선 선택',
        'Process → Export → Navigation',
        'Before 또는 After 선택',
        '출력 폴더 지정 후 Export',
    ]
    for i, s in enumerate(caris_steps, 1):
        doc.add_paragraph(f'{i}) {s}')

    doc.add_paragraph()
    doc.add_heading('지원 Nav 파일 형식 (6종 자동 감지)', level=3)
    add_styled_table(doc,
        ['형식', '설명', '예시'],
        [
            ['형식 1', 'After 표준', 'YYYY-DDD HH:MM:SS.sss Lat Lon Depth Hex'],
            ['형식 2', 'Before 표준', 'YYYY-DDD HH:MM:SS.sss Lat Lon'],
            ['형식 3', '날짜-시간', 'YYYY/MM/DD HH:MM:SS Lat Lon Depth'],
            ['형식 4', '탭 구분', 'YYYY-DDD\\tHH:MM:SS.sss\\tLat\\tLon\\t...'],
            ['형식 5', '콤마 구분', 'YYYY-DDD,HH:MM:SS.sss,Lat,Lon,Depth'],
            ['형식 6', '확장 형식', '공백/탭 혼합 등'],
        ],
        col_widths=[2.5, 3.5, 10]
    )

    doc.add_page_break()

    # 3.3
    doc.add_heading('3.3  기준항정보.txt', level=2)
    doc.add_paragraph(
        '기준항 목록과 조화상수를 정의하는 탭(\\t) 구분 텍스트 파일입니다. '
        'info/ 폴더에서 자동 탐색되거나 GUI에서 수동 지정할 수 있습니다.'
    )
    doc.add_paragraph()
    doc.add_heading('필드 구성 (17필드)', level=3)
    add_styled_table(doc,
        ['#', '필드명', '설명', '단위'],
        [
            ['1', 'Use', '사용 여부', 'T/F'],
            ['2', 'Seq', '순번', '정수'],
            ['3', 'Name', '기준항 명칭', '문자열'],
            ['4', 'Lon', '경도', 'DD'],
            ['5', 'Lat', '위도', 'DD'],
            ['6', 'M2Amp', 'M2 분조 진폭', 'cm'],
            ['7', 'M2Phase', 'M2 분조 위상', 'degree'],
            ['8', 'S2Amp', 'S2 분조 진폭', 'cm'],
            ['9', 'S2Phase', 'S2 분조 위상', 'degree'],
            ['10', 'K1Amp', 'K1 분조 진폭', 'cm'],
            ['11', 'K1Phase', 'K1 분조 위상', 'degree'],
            ['12', 'O1Amp', 'O1 분조 진폭', 'cm'],
            ['13', 'O1Phase', 'O1 분조 위상', 'degree'],
            ['14', 'SprRange', '대조차', 'cm'],
            ['15', 'SprRise', '대조 상승량', 'cm'],
            ['16', 'MSL', '평균해면', 'cm'],
            ['17', 'MHWI', '평균고조간격', 'hour'],
        ],
        col_widths=[1.2, 2.5, 5, 2.5]
    )

    add_info_box(doc, '⚠ 중요', [
        'Name 필드는 TOPS 조위 파일 관측소명과 정확히 일치해야 합니다.',
    ], 'FDE8E8')

    # 3.4
    doc.add_heading('3.4  표준개정수 DB (Co-tidal Database)', level=2)
    doc.add_paragraph(
        '해역 전체를 0.5°×0.5° 격자로 분할하고, 각 격자점의 조석 파라미터를 저장합니다.'
    )
    doc.add_paragraph()
    doc.add_heading('구성', level=3)
    doc.add_paragraph('• File_Catalog.txt — 격자 섹터별 NetCDF 파일 인덱스')
    doc.add_paragraph('• CT/ 폴더 — 실제 .nc 격자 파일')

    doc.add_heading('NetCDF 변수', level=3)
    add_styled_table(doc,
        ['변수명', '설명', '단위'],
        [
            ['SprRange', '대조차', 'cm'],
            ['DL_MSL', '기본수준면-평균해면 차', 'cm'],
            ['MHWI', '평균고조간격', 'hour'],
        ],
        col_widths=[4, 8, 4]
    )
    doc.add_page_break()


def add_section_4(doc):
    """4. 사용 방법."""
    doc.add_heading('4. 사용 방법', level=1)

    # 4.1 GUI
    doc.add_heading('4.1  GUI 모드', level=2)
    doc.add_heading('실행 방법', level=3)
    doc.add_paragraph('• EXE: TideBedPy.exe 더블클릭')
    doc.add_paragraph('• Python: python gui.py')
    doc.add_paragraph('• 배치파일: TideBedPy_GUI.bat')
    doc.add_paragraph()

    # (1) 입력 파일
    add_info_box(doc, '(1) 입력 파일 설정', [
        '항적 폴더 — Nav 파일 폴더 선택 ([찾아보기] 또는 직접 입력)',
        '조위 폴더 — TOPS 조위 파일 폴더 선택',
        '출력 파일 — 결과 .tid 파일 저장 경로 지정',
    ], 'E8F0FE')

    add_info_box(doc, '(2) 개정수 DB 설정', [
        '개정수 DB 경로 — File_Catalog.txt 폴더 (info/ 자동 탐색)',
        '기준항정보 — 기준항정보.txt 파일 (자동 탐색 가능)',
    ], 'E8F0FE')

    # (3) 보정 옵션
    doc.add_heading('(3) 보정 옵션', level=3)
    add_styled_table(doc,
        ['옵션', '설명', '기본값'],
        [
            ['조위 유형', '실측조위 / 예측조위 (실측 권장)', '실측조위'],
            ['기준항 적용 개수', 'IDW 가중평균 시 사용할 상위 기준항 수 (1~10)', '10'],
            ['기준 시간대', '27개 UTC 오프셋 중 선택 (GMT, KST, JST, CST 등)', 'GMT'],
            ['출력 시간 간격', '.tid 출력 간격 (초)', '원본 간격 유지'],
            ['허용 편차', '검증 그래프 오차 허용 범위 (cm)', '1.0'],
            ['상세 출력', '체크 시 .tid.detail 생성', '미체크'],
            ['그래프 출력', '체크 시 .tid.png, .tid.map.png 등 생성', '미체크'],
        ],
        col_widths=[4, 8, 4]
    )

    doc.add_paragraph()
    add_info_box(doc, '(4) 실행', [
        '[보정 수행] 클릭 → 실시간 진행률 + 로그 + 결과 요약 표시',
    ], 'E8F8E8')

    add_info_box(doc, '(5) 세팅 관리', [
        '[세팅 저장] — 현재 설정을 프리셋 파일로 저장',
        '[세팅 불러오기] — 이전 프리셋 불러오기',
        '[INI 불러오기] — TideBedLite .ini 호환 모드',
    ], 'E8F0FE')

    doc.add_page_break()

    # 4.2 CLI
    doc.add_heading('4.2  CLI 모드', level=2)
    doc.add_heading('기본 사용법', level=3)

    add_info_box(doc, '명령어', [
        'python main.py --nav <항적폴더> --tide <조위폴더> -o <출력파일>',
    ], 'F0F0F0')

    doc.add_heading('사용 예시', level=3)
    add_info_box(doc, '예시 1 — 기본 실행', [
        'python main.py --nav Navi/After --tide 실측조위 -o result.tid',
    ], 'F0F0F0')
    add_info_box(doc, '예시 2 — INI 파일 사용', [
        'python main.py --ini setting/TideBedLite.ini -o result.tid',
    ], 'F0F0F0')
    add_info_box(doc, '예시 3 — 전체 옵션', [
        'python main.py --nav Navi/After --tide 실측조위 --db info \\',
        '    --stations info/기준항정보.txt -o result.tid \\',
        '    --type observed --rank-limit 5 --time-interval 10 \\',
        '    --detail --kst -v',
    ], 'F0F0F0')

    doc.add_paragraph()
    doc.add_heading('옵션 목록', level=3)
    add_styled_table(doc,
        ['옵션', '설명', '기본값'],
        [
            ['--nav <경로>', '항적 폴더', '(필수)'],
            ['--tide <경로>', '조위 폴더', '(필수)'],
            ['--db <경로>', '개정수 DB 폴더', 'info/'],
            ['--stations <경로>', '기준항정보.txt', '자동탐색'],
            ['-o <경로>', '출력 .tid 파일', '(필수)'],
            ['--type <유형>', 'observed / predicted', 'observed'],
            ['--rank-limit <N>', '기준항 적용 개수 (1~10)', '10'],
            ['--time-interval <초>', '출력 시간 간격', '원본 유지'],
            ['--detail', '.tid.detail 생성', '미생성'],
            ['--kst', '시간대 KST(UTC+9)', 'GMT'],
            ['--validate <참조tid>', '참조 .tid와 비교 검증', '미실행'],
            ['-v, --verbose', '상세 로그', '비활성'],
        ],
        col_widths=[4.5, 7, 4.5]
    )

    doc.add_page_break()

    # 4.3 검증 모드
    doc.add_heading('4.3  검증 모드', level=2)
    doc.add_paragraph(
        '기존 TideBedLite 결과와 비교하여 정확성을 검증합니다.'
    )
    add_info_box(doc, '사용 명령어', [
        'python main.py --nav Navi --tide 조위 -o test.tid --validate ref.tid',
    ], 'F0F0F0')

    doc.add_paragraph()
    doc.add_paragraph('검증 기준: ±0.01m (1cm)')
    doc.add_paragraph()
    doc.add_heading('출력', level=3)
    doc.add_paragraph('• 콘솔에 PASS/FAIL + 통계 요약')
    doc.add_paragraph('• .tid.compare.png: 비교 시계열 그래프')
    doc.add_paragraph('• 최대 오차, 평균 오차, RMSE 표시')
    doc.add_page_break()


def add_section_5(doc):
    """5. 출력 파일 설명."""
    doc.add_heading('5. 출력 파일 설명', level=1)

    add_styled_table(doc,
        ['파일', '설명'],
        [
            ['result.tid', 'TOPS 호환 조석보정 결과 (CARIS Import용)\n형식: YYYY/MM/DD HH:MM:SS  Tc_m  0.0'],
            ['result.tid.detail', '상세 보정 정보 (--detail 옵션)\n기준항별 가중치, HRatio, TimeCorrector 등'],
            ['result.tid.err', '에러 포인트 목록\n매칭 실패, 격자 범위 초과 등'],
            ['result.tid.png', '조석보정 시계열 그래프\nX축: 시간, Y축: Tc (m)'],
            ['result.tid.compare.png', '참조값 비교 그래프 (검증 시)\n오차 분포 + 통계 표시'],
            ['result.tid.map.png', '기준항 위치 + 항적 경로 지도\n사용/미사용 기준항 구분 표시'],
            ['result.tid.corrmap.png', '보정결과 컬러맵 지도\nTc 값을 색상으로 매핑'],
        ],
        col_widths=[5, 11]
    )

    doc.add_paragraph()
    doc.add_heading('.tid 파일 예시', level=2)
    add_info_box(doc, '.tid 출력 형식', [
        '2024/01/15 09:30:00   3.45  0.0',
        '2024/01/15 09:30:10   3.44  0.0',
        '2024/01/15 09:30:20   3.43  0.0',
    ], 'F0F0F0')
    doc.add_page_break()


def add_section_6(doc):
    """6. 핵심 알고리즘 및 공식."""
    doc.add_heading('6. 핵심 알고리즘 및 공식', level=1)

    # 6.1
    doc.add_heading('6.1  Vincenty 측지 거리 계산', level=2)
    doc.add_paragraph(
        '항적 포인트와 기준항 사이의 정확한 지구 표면 거리를 계산합니다. '
        'WGS84 타원체 모델 기반 Karney 알고리즘 (geographiclib)을 사용합니다.'
    )
    doc.add_paragraph()
    doc.add_heading('WGS84 파라미터', level=3)
    add_styled_table(doc,
        ['파라미터', '값', '설명'],
        [
            ['a', '6,378,137.0 m', '장반경'],
            ['f', '1/298.257223563', '편평률'],
            ['b', '6,356,752.314245 m', '단반경'],
        ],
        col_widths=[3, 5, 5]
    )
    doc.add_paragraph()
    doc.add_paragraph('입력: (lon₁, lat₁), (lon₂, lat₂)  →  출력: 측지 거리 d (m)')
    doc.add_paragraph('정밀도: ~0.5mm 이하')

    # 6.2
    doc.add_heading('6.2  IDW (Inverse Distance Weighting) 가중평균', level=2)
    doc.add_paragraph('역거리 가중법으로 가까운 기준항에 더 큰 가중치를 부여합니다.')
    doc.add_paragraph()

    add_info_box(doc, 'IDW 계산 과정', [
        'Step 1 — 거리 계산',
        '    d_i = VincentyDistance(P, S_i)',
        '',
        'Step 2 — 원시 가중치',
        '    w_i = 1 / d_i²',
        '',
        'Step 3 — 전체 기준항 대상 정규화',
        '    W_i = w_i / Σ(w_j)    ← j = 모든 기준항',
        '',
        'Step 4 — 상위 RankLimit개 선택 (거리순)',
        '',
        'Step 5 — 재정규화 + 가중평균',
        '    Tc = Σ(EstimHeight_i × W_i) / Σ(W_i)',
        '         (i = 상위 RankLimit개)',
    ], 'E8F0FE')

    add_info_box(doc, '⚠ 참고', [
        'Step 3 정규화는 전체 기준항 대상, Step 5는 상위 N개만 대상 (원본 동일)',
    ], 'FFF8E1')

    # 6.3
    doc.add_heading('6.3  Co-tidal 격자 보간 (이중선형 보간)', level=2)
    doc.add_paragraph(
        '항적 위치에서 Co-tidal 격자의 SprRange, DL_MSL, MHWI를 '
        '이중선형 보간(Bilinear Interpolation)으로 구합니다.'
    )
    doc.add_paragraph('격자 크기: 0.5°×0.5° (각 격자점에 조석 파라미터 저장)')
    doc.add_paragraph()

    add_info_box(doc, '이중선형 보간 공식', [
        'V00 ──── V10        dx = (lon - lon₀) / (lon₁ - lon₀)',
        ' |    P    |        dy = (lat - lat₀) / (lat₁ - lat₀)',
        'V01 ──── V11',
        '',
        'V = V00·(1-dx)·(1-dy) + V10·dx·(1-dy)',
        '  + V01·(1-dx)·dy     + V11·dx·dy',
    ], 'E8F0FE')

    doc.add_paragraph()
    doc.add_heading('보간 대상', level=3)
    add_styled_table(doc,
        ['변수', '설명', '단위'],
        [
            ['SprRange_nav', '항적 위치의 대조차', 'cm'],
            ['MSL_nav', '항적 위치의 평균해면', 'cm'],
            ['MHWI_nav', '항적 위치의 평균고조간격', 'hour'],
        ],
        col_widths=[5, 7, 4]
    )

    doc.add_page_break()

    # 6.4
    doc.add_heading('6.4  Akima 스플라인 보간', level=2)
    doc.add_paragraph(
        'TOPS 조위 데이터(10분~1시간 간격)를 1분 간격으로 리샘플링합니다.'
    )
    doc.add_paragraph('라이브러리: scipy.interpolate.Akima1DInterpolator')
    doc.add_paragraph()
    doc.add_heading('처리 과정', level=3)
    steps = [
        'TOPS 데이터 → 시간-수위 쌍 로드',
        'Akima 보간기 생성',
        '1분 간격 리샘플링',
        '항적 시각과 이진탐색 매칭 (±2분 이내)',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'{i}) {s}')

    doc.add_paragraph()
    doc.add_heading('Akima 스플라인의 장점', level=3)
    doc.add_paragraph('• 국소적 보간 (한 구간 변화가 먼 구간에 미영향)')
    doc.add_paragraph('• 과도한 진동 없이 부드러운 보간')
    doc.add_paragraph('• 조석 데이터의 물리적 특성에 적합')

    # 6.5
    doc.add_heading('6.5  조석보정 계산 단계', level=2)
    doc.add_paragraph('각 항적 포인트에 대해 순차 수행:')
    doc.add_paragraph()

    add_info_box(doc, 'Step 1 — IDW 가중치 계산', [
        '항적↔기준항 측지 거리 계산 → 역거리 제곱 가중치 정규화',
    ], 'E8F0FE')

    add_info_box(doc, 'Step 2 — Co-tidal 값 보간', [
        '항적 위치에서 이중선형 보간',
        '→ SprRange_nav, MSL_nav, MHWI_nav',
    ], 'E8F0FE')

    add_info_box(doc, 'Step 3 — 각 기준항별 추정 수위', [
        '(a) HRatio_i = SprRange_nav / SprRange_station_i',
        '',
        '(b) TimeCorrector_i = (MHWI_nav - MHWI_station_i) × 2.0',
        '',
        '(c) correctedTime_i = obsTime + (9.0 - utc_offset) - TimeCorrector_i',
        '',
        '(d) orgHeight_i = find_level(station_series_i, correctedTime_i)',
        '',
        '(e) EstimHeight_i = (orgHeight_i - MSL_st_i) × HRatio_i + MSL_nav',
    ], 'E8F0FE')

    add_info_box(doc, 'Step 4 — 가중평균 → 최종 Tc', [
        'Tc(cm) = Σ(EstimHeight_i × W_i) / Σ(W_i)',
        '          (i = 상위 RankLimit개)',
    ], 'E8F0FE')

    add_info_box(doc, 'Step 5 — 단위 변환 + 출력', [
        'Tc(m) = Tc(cm) / 100.0  →  소수점 2자리 반올림',
        '.tid: "YYYY/MM/DD HH:MM:SS  Tc_m  0.0"',
    ], 'E8F8E8')

    # 6.6
    doc.add_heading('6.6  시간대 처리', level=2)
    doc.add_paragraph(
        '조위 데이터(TOPS)는 일반적으로 KST(UTC+9) 기준입니다. '
        '항적 시간대에 따른 변환이 필요합니다.'
    )
    doc.add_paragraph()
    doc.add_heading('공식', level=3)
    add_info_box(doc, '시간 변환', [
        'time_shift = 9.0 - utc_offset',
        'correctedTime = obsTime + time_shift - TimeCorrector',
    ], 'E8F0FE')

    doc.add_paragraph()
    doc.add_heading('예시', level=3)
    add_styled_table(doc,
        ['Nav 시간대', 'utc_offset', 'time_shift', '설명'],
        [
            ['GMT (UTC+0)', '0.0', '+9.0h', '9시간 더해서 KST 매칭'],
            ['KST (UTC+9)', '9.0', '0.0h', '동일 시간대, 보정 없음'],
            ['CST (UTC+8)', '8.0', '+1.0h', '1시간 더해서 KST 매칭'],
        ],
        col_widths=[4, 3, 3, 6]
    )
    doc.add_page_break()


def add_section_7(doc):
    """7. 시간대 설정 안내."""
    doc.add_heading('7. 시간대 설정 안내', level=1)

    doc.add_paragraph('지원 시간대 (27종):')
    doc.add_paragraph()

    add_styled_table(doc,
        ['표기', 'UTC 오프셋', '설명'],
        [
            ['GMT', '+0', '그리니치 표준시 (= UTC)'],
            ['UTC+1 ~ +7', '+1 ~ +7', '유럽, 중동, 동남아 등'],
            ['CST', '+8', '중국 표준시'],
            ['KST', '+9', '한국 표준시'],
            ['JST', '+9', '일본 표준시'],
            ['UTC+10 ~ +12', '+10 ~ +12', '호주, 피지 등'],
            ['UTC-1 ~ -12', '-1 ~ -12', '대서양, 미주 등'],
            ['UTC+5:30', '+5.5', '인도 표준시'],
        ],
        col_widths=[4, 3, 9]
    )

    doc.add_paragraph()
    doc.add_heading('선택 가이드', level=2)
    doc.add_paragraph('• 대부분의 한국 수로측량 → KST 또는 GMT')
    doc.add_paragraph('• CARIS HIPS Export 시 시간대 확인 후 동일 설정')
    doc.add_paragraph('• 잘못된 시간대 = 조석보정 오류의 가장 흔한 원인')
    doc.add_page_break()


def add_section_8(doc):
    """8. TideBedLite 대비 개선점."""
    doc.add_heading('8. TideBedLite 대비 개선점', level=1)

    add_styled_table(doc,
        ['#', '항목', '원본 (.NET)', 'TideBedPy (Python)'],
        [
            ['1', 'Nav 포맷 지원', '2가지', '6가지 자동 감지'],
            ['2', '시간대', 'KST/GMT만', '27개 UTC 오프셋'],
            ['3', '조위 보간', '원시 직접 탐색', 'Akima 스플라인'],
            ['4', '진행률 표시', '없음', '실시간 프로그레스'],
            ['5', '시각화', '없음', '그래프 + 지도 4종'],
            ['6', '설정 관리', 'INI 파일만', '프리셋 저장/불러오기'],
            ['7', '인터페이스', 'GUI만', 'CLI + GUI'],
            ['8', '배포', '.NET 런타임 필요', 'EXE 독립 실행'],
            ['9', '인코딩', '수동 지정', '자동 감지'],
            ['10', '검증', '없음', '참조 .tid 비교'],
        ],
        col_widths=[1.2, 3.5, 4.5, 4.5]
    )
    doc.add_page_break()


def add_section_9(doc):
    """9. 문제 해결 (FAQ)."""
    doc.add_heading('9. 문제 해결 (FAQ)', level=1)

    faqs = [
        (
            'Q1. "매칭된 기준항이 없습니다"',
            '원인: 기준항정보.txt Name ≠ TOPS 파일 관측소명',
            [
                '기준항정보.txt Name 필드 확인',
                'TOPS 파일 1행 관측소명 확인',
                '두 값이 정확히 동일한지 비교 (공백, 특수문자 주의)',
                'Use 필드가 TRUE인지 확인',
            ]
        ),
        (
            'Q2. "항적 데이터를 불러올 수 없습니다"',
            '원인: Nav 파일이 지원 6가지 형식에 해당하지 않음',
            [
                'Nav 파일을 텍스트 편집기로 열어 형식 확인',
                '빈 행, 잘못된 구분자, 깨진 문자 확인',
                'CARIS에서 Navigation Export 재수행',
            ]
        ),
        (
            'Q3. "Co-tidal 격자 로드 실패"',
            '원인: 표준개정수 DB 파일을 찾을 수 없거나 손상',
            [
                'File_Catalog.txt 존재 확인',
                'CT/ 폴더와 .nc 파일 존재 확인',
                'GUI에서 경로 수동 지정',
                'netCDF4 설치 확인: pip install netCDF4',
            ]
        ),
        (
            'Q4. 인코딩 에러',
            '원인: 파일 인코딩 문제',
            [
                '파일을 UTF-8(BOM 없음)로 저장 권장',
                'EUC-KR도 자동 감지됨',
                'chardet 설치 확인: pip install chardet',
            ]
        ),
        (
            'Q5. 조석보정값 비정상',
            '원인: 시간대 설정 오류 (가장 흔한 원인)',
            [
                'Nav 파일 시간대 확인 (GMT/KST)',
                '프로그램 시간대 설정 일치 확인',
                '--detail로 기준항별 파라미터 검토',
            ]
        ),
        (
            'Q6. 에러 포인트가 많음',
            '원인: 조위 데이터 시간 범위 < 항적 시간 범위',
            [
                '조위/항적 시작·종료 시각 확인',
                '조위 데이터가 항적 전체 기간을 커버하는지 확인',
                '바다누리에서 추가 조위 다운로드',
            ]
        ),
        (
            'Q7. GUI 실행 불가',
            '원인: 환경 설정 문제',
            [
                'Python 3.8+ 확인: python --version',
                '의존성 설치: pip install numpy scipy netCDF4 ...',
                'tkinter 포함 Python인지 확인',
                'EXE: 안티바이러스 차단 여부 확인',
            ]
        ),
    ]

    for q_title, cause, solutions in faqs:
        doc.add_heading(q_title, level=2)
        p_cause = doc.add_paragraph()
        run_c = p_cause.add_run(cause)
        run_c.font.name = FONT_KR
        run_c._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
        run_c.font.size = Pt(10)
        run_c.font.color.rgb = MID_GRAY
        run_c.italic = True

        doc.add_paragraph()
        p_sol = doc.add_paragraph()
        run_sol = p_sol.add_run('해결:')
        run_sol.font.name = FONT_KR
        run_sol._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
        run_sol.bold = True

        for i, s in enumerate(solutions, 1):
            doc.add_paragraph(f'  {i}. {s}')

        doc.add_paragraph()

    doc.add_page_break()


def add_section_10(doc):
    """10. 라이선스 및 저작권."""
    doc.add_heading('10. 라이선스 및 저작권', level=1)

    doc.add_heading('원본 프로그램', level=2)
    doc.add_paragraph('TideBedLite v1.05')
    doc.add_paragraph('Copyright (c) 2014, KHOA (국립해양조사원) / GeoSR Inc.')
    doc.add_paragraph('Co-tidal 격자 기반 조석보정 .NET 프로그램')

    doc.add_heading('Python 재구현', level=2)
    doc.add_paragraph('TideBedPy v2.3.0')
    doc.add_paragraph('재작성: Junhyub, 2025')
    doc.add_paragraph('원본 알고리즘 충실히 재현 + 추가 기능 개선')

    doc.add_heading('사용 라이브러리', level=2)
    add_styled_table(doc,
        ['라이브러리', '라이선스', '용도'],
        [
            ['numpy', 'BSD-3-Clause', '수치 연산'],
            ['scipy', 'BSD-3-Clause', 'Akima 스플라인'],
            ['netCDF4', 'MIT', 'Co-tidal 격자 읽기'],
            ['geographiclib', 'MIT', '측지 거리 계산'],
            ['chardet', 'LGPL-2.1', '인코딩 자동 감지'],
            ['matplotlib', 'PSF (BSD-like)', '그래프 및 지도 시각화'],
        ],
        col_widths=[4, 4, 8]
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # 푸터
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.color.rgb = BLUE
    run.font.size = Pt(10)

    for text in [
        'TideBedPy v2.3.0',
        'Original: TideBedLite v1.05 (c) 2014, KHOA / GeoSR Inc.',
        'Python: Junhyub, 2025',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = FONT_KR
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KR)
        run.font.size = Pt(10)
        run.font.color.rgb = MID_GRAY


def set_page_margins(doc):
    """페이지 여백 설정."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def main():
    print('  TideBedPy Word Manual Builder')
    print('  ' + '=' * 50)

    doc = Document()

    # 기본 설정
    setup_styles(doc)
    set_page_margins(doc)

    # 내용 생성
    print('  [1/10] 표지 생성...')
    add_cover(doc)

    print('  [2/10] 목차 생성...')
    add_toc(doc)

    print('  [3/10] 1. 프로그램 개요...')
    add_section_1(doc)

    print('  [4/10] 2. 시스템 요구사항...')
    add_section_2(doc)

    print('  [5/10] 3. 데이터 준비...')
    add_section_3(doc)

    print('  [6/10] 4. 사용 방법...')
    add_section_4(doc)

    print('  [7/10] 5. 출력 파일 설명...')
    add_section_5(doc)

    print('  [8/10] 6. 핵심 알고리즘...')
    add_section_6(doc)

    print('  [9/10] 7~9. 시간대/개선점/FAQ...')
    add_section_7(doc)
    add_section_8(doc)
    add_section_9(doc)

    print('  [10/10] 10. 라이선스...')
    add_section_10(doc)

    # 저장
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    size_kb = os.path.getsize(OUT_PATH) / 1024
    print(f'\n  [OK] 매뉴얼 생성 완료: {OUT_PATH}')
    print(f'       크기: {size_kb:.1f} KB')


if __name__ == '__main__':
    main()
